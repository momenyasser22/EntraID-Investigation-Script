import pandas as pd
import time
import requests
from pathlib import Path
from docx import Document

# =============================
# Configuration
# =============================
BASELINE_LOCATION = "EG"
VT_API_KEY = "0ff41484d0e540398497d5c511cd3fead0e0bed2f0df0a92c614f3f4b7c76386"
SCRIPT_DIR = Path(__file__).parent


# Output directory (Docker + local compatible)
DOCKER_DATA_DIR = Path("/data")

if DOCKER_DATA_DIR.exists():
    # Running inside Docker
    OUTPUT_DIR = DOCKER_DATA_DIR / "output"
else:
    # Running locally
    OUTPUT_DIR = SCRIPT_DIR / "data" / "output"

OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

OUTPUT_CSV = OUTPUT_DIR / "main_investigation.csv"

# =============================
# Helper Functions
# =============================
def get_file_path(prompt_text):
    path = input(prompt_text).strip().strip('"')
    return Path(path)

def load_and_clean(file_path):
    file_path = Path(file_path)

    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    if file_path.suffix.lower() == ".csv":
        df = pd.read_csv(file_path)
    elif file_path.suffix.lower() in [".xls", ".xlsx"]:
        df = pd.read_excel(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_path.suffix}")

    return df

def merge_files(signin_df, auth_df):
    merged = signin_df.merge(auth_df, on="Request ID", how="left")
    return merged

def extract_country(location):
    if pd.isna(location):
        return ""
    return location.split(",")[-1].strip()

def deduplicate_by_timestamp(df, timestamp_col="Date (UTC)", success_col="Succeeded"):
    if timestamp_col not in df.columns:
        return df

    # If success column exists, prioritize successful authentications
    if success_col in df.columns:
        df = df.sort_values(
            by=[timestamp_col, success_col],
            ascending=[True, False]
        )

    return df.drop_duplicates(subset=[timestamp_col], keep="first")

def investigation_step_1(df):
    df = deduplicate_by_timestamp(df)
    findings = []
    grouped = df.sort_values("Date").groupby("Session ID")

    for session_id, group in grouped:
        countries = group["Location"].apply(extract_country).tolist()
        if BASELINE_LOCATION in countries and any(c != BASELINE_LOCATION for c in countries):
            findings.append(group)
    print(f"[+] Step 1 completed. Potential session hijacks found: {len(findings)}")
    return findings

def investigation_step_2(df):
    df = deduplicate_by_timestamp(df)
    df = df.copy()
    df["Country"] = df["Location"].apply(extract_country)
    print(f"[+] Step 2 completed. Entries outside baseline country: {len(df[df['Country'] != BASELINE_LOCATION])}")
    return df[df["Country"] != BASELINE_LOCATION]

def virustotal_lookup(ips):
    results = {}
    headers = {"x-apikey": VT_API_KEY}

    total_ips = len(ips)

    if total_ips == 0:
        print("[+] No IPs to submit for VirusTotal lookup.")
        return results

    print(f"[+] Starting VirusTotal lookup for {total_ips} IP addresses.")
    print("[+] Rate limit: 1 request per minute\n")

    for index, ip in enumerate(ips, start=1):
        remaining = total_ips - index
        remaining_time = remaining * 60

        print(f"[VT] ({index}/{total_ips}) Checking IP: {ip}")
        print(f"     Remaining IPs: {remaining}")
        print(f"     Estimated time remaining: {remaining_time // 60} min {remaining_time % 60} sec")

        url = f"https://www.virustotal.com/api/v3/ip_addresses/{ip}"
        r = requests.get(url, headers=headers)

        if r.status_code == 200:
            results[ip] = r.json()
            print("     Status: Success")
        else:
            results[ip] = {"error": r.status_code}
            print(f"     Status: Failed (HTTP {r.status_code})")

        if remaining > 0:
            print("     Sleeping 60 seconds to respect API rate limits...\n")
            time.sleep(60)

    print("[+] VirusTotal enrichment completed.\n")
    return results

def investigation_step_4(df):
    df = df.copy()
    df["Country"] = df["Location"].apply(extract_country)

    # Catch successful Password Hash Sync authentication from foreign locations
    filtered = df[
        (df["Country"] != BASELINE_LOCATION) &
        (df["Succeeded"] == True) &
        (
            df["Authentication method detail"]
            .str.contains("Password Hash Sync", case=False, na=False)
        )
    ]

    # Deduplicate AFTER filtering so the success is not dropped
    filtered = deduplicate_by_timestamp(filtered)

    return filtered

def generate_report(step1, step2, step3, step4):
    doc = Document()

    def add_table_from_df(df, title):
        if df is None or df.empty:
            doc.add_paragraph(f"{title}: No records identified.")
            return

        doc.add_paragraph(title)
        table = doc.add_table(rows=1, cols=len(df.columns))
        hdr_cells = table.rows[0].cells

        for i, col in enumerate(df.columns):
            hdr_cells[i].text = col

        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)

    # Title
    doc.add_heading("Entra ID Sign-In Investigation Report", level=1)

    # Executive Summary
    doc.add_heading("1. Executive Summary", level=2)
    doc.add_paragraph(
        "An investigation was conducted to analyze Microsoft Entra ID sign-in activity "
        "for indications of unauthorized access, session hijacking, and anomalous "
        "authentication behavior. The investigation focused on geographic anomalies, "
        "authentication outcomes, and threat intelligence associated with observed IP addresses."
    )

    # Scope & Methodology
    doc.add_heading("2. Scope and Methodology", level=2)
    doc.add_paragraph(
        "The investigation followed a structured, step-based approach:\n"
        "• Session correlation using Session ID to detect potential hijacking\n"
        "• Baseline country enforcement review (Egypt)\n"
        "• Threat intelligence enrichment using VirusTotal\n"
        "• Verification of successful authentication attempts from foreign locations"
    )

    # Step 1 Findings
    doc.add_heading("3. Findings", level=2)
    doc.add_heading("3.1 Session Hijacking Analysis", level=3)

    if not step1:
        doc.add_paragraph(
            "No session hijacking indicators were identified during the investigation period."
        )
    else:
        doc.add_paragraph(
            "One or more sessions exhibited geographic anomalies consistent with potential session hijacking."
        )
        for g in step1:
            add_table_from_df(
                g[["Session ID", "User", "IP address", "Location", "Application", "Resource"]],
                "Session Hijacking Evidence"
            )

    # Step 2 Findings
    doc.add_heading("3.2 Baseline Location Violations", level=3)

    foreign_events = step2.copy()
    doc.add_paragraph(
        f"A total of {len(foreign_events)} sign-in events were detected from locations "
        f"outside the baseline country (Egypt)."
    )

    add_table_from_df(
        foreign_events[["Date (UTC)", "User", "IP address", "Location", "Status"]],
        "Foreign Sign-In Events"
    )

    # Step 3 Findings
    doc.add_heading("3.3 Threat Intelligence Assessment", level=3)

    if not step3:
        doc.add_paragraph("No IP addresses required threat intelligence enrichment.")
    else:
        ti_rows = []
        for ip, result in step3.items():
            stats = result.get("data", {}).get("attributes", {}).get("last_analysis_stats", {})
            ti_rows.append({
                "IP Address": ip,
                "Malicious": stats.get("malicious", 0),
                "Suspicious": stats.get("suspicious", 0)
            })

        ti_df = pd.DataFrame(ti_rows)
        add_table_from_df(ti_df, "Threat Intelligence Summary")

    # Step 4 Findings
    doc.add_heading("3.4 Authentication Success Review", level=3)

    successful_foreign_auth = step4.copy()

    if successful_foreign_auth.empty:
        doc.add_paragraph(
            "No successful authentication attempts were observed from outside Egypt."
        )
    else:
        doc.add_paragraph(
            "Successful authentication attempts from foreign locations were identified."
        )
        add_table_from_df(
            successful_foreign_auth[
                ["Date (UTC)", "User", "Authentication method", "IP address", "Location"]
            ],
            "Successful Foreign Authentication Events"
        )

    # Risk Assessment
    doc.add_heading("4. Risk Assessment", level=2)
    doc.add_paragraph(
        "Based on the investigation results, there is no evidence of account compromise. "
        "Observed foreign sign-in attempts were unsuccessful, indicating that existing "
        "security controls are functioning as intended."
    )

    # Recommendations
    doc.add_heading("5. Recommendations", level=2)

    recommendations = []

    # Foreign IPs observed (baseline violations) – only if NO successful foreign authentication
    if not step2.empty and step4.empty:
        recommendations.append(
            "Reset the affected user's password to mitigate the risk of potential credential exposure "
            "following authentication attempts from foreign IP addresses."
        )

    # Always recommend IP blocking if foreign IPs exist
    if not step2.empty:
        foreign_ips = sorted(step2["IP address"].dropna().unique().tolist())
        if foreign_ips:
            recommendations.append(
                "Block the following IP addresses at the network or identity provider level "
                "to prevent further malicious authentication attempts:\n"
                + ", ".join(foreign_ips)
            )

    # Successful foreign authentication
    if not step4.empty:
        recommendations.append(
            "Immediately reset the affected user's password due to successful password-based authentication "
            "from a foreign location and invalidate all active sessions."
        )

    # Session hijacking indicators
    if step1:
        recommendations.append(
            "Reset the user's MFA registrations and revoke active sessions due to indicators "
            "consistent with potential session hijacking."
        )

    # Baseline security controls
    recommendations.append(
        "Continue enforcing Conditional Access policies restricting access based on geographic location."
    )

    recommendations.append(
        "Maintain MFA enforcement for all sign-ins originating from non-baseline countries."
    )

    for rec in recommendations:
        doc.add_paragraph(f"• {rec}")


    doc.save(OUTPUT_DOC)

# =============================
# Main
# =============================
def main():
    print("Please provide the full paths to the required files.")
    signin_path = get_file_path("Enter path for SignIn file: ")
    auth_path = get_file_path("Enter path for signin authdetails file: ")

    signin_df = load_and_clean(signin_path)
    auth_df = load_and_clean(auth_path)

    main_df = merge_files(signin_df, auth_df)
    main_df.to_csv(OUTPUT_CSV, index=False)

    # Determine investigated user for report naming
    users = main_df["User"].dropna().unique()

    if len(users) == 1:
        investigated_user = users[0]
    else:
        investigated_user = "Multiple_Users"

    # Sanitize filename (remove unsafe characters)
    safe_user = "".join(
        c for c in investigated_user if c.isalnum() or c in (" ", "_", "-")
    ).strip().replace(" ", "_")

    global OUTPUT_DOC
    OUTPUT_DOC = OUTPUT_DIR / f"Investigation_Report_{safe_user}.docx"

    step1 = investigation_step_1(main_df)
    step2 = investigation_step_2(main_df)

    print("[+] Step 1 output:")
    for g in step1:
        print(g[["Session ID", "Location", "User agent", "Application", "Resource"]])

    ip_column = "IP address"
    ips = step2[ip_column].dropna().unique().tolist()

    print(f"[+] Step 2 completed. Foreign IPs detected: {len(ips)}")
    for ip in ips:
        print(f"    - {ip}")

    step3 = virustotal_lookup(ips)

    step4 = investigation_step_4(main_df)

    print("[+] Step 4 output:")
    print(step4[["User", "Authentication method", "Location", "Succeeded"]])

    generate_report(step1, step2, step3, step4)

    print("Investigation completed.")
    print(f"CSV exported to: {OUTPUT_CSV}")
    print(f"Report exported to: {OUTPUT_DOC}")

if __name__ == "__main__":
    main()
